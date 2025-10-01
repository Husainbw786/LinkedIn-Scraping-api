import io
import os
import json
import hashlib
from typing import Dict, Any, List, Optional
import pdfplumber
import PyPDF2
from docx import Document
from openai import OpenAI
from utils.logger_config import setup_logging

logger = setup_logging()
class ResumeEmbeddingService:
    """Service for processing resumes and generating embeddings"""
    
    def __init__(self):
        self.openai_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        self.embedding_model = "text-embedding-ada-002"
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """
        Extract text from PDF content using multiple methods for better accuracy
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Extracted text string
        """
        try:
            text = ""
            
            # Method 1: Try pdfplumber first (better for complex layouts)
            try:
                with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    logger.info("✅ Text extracted using pdfplumber")
                    return self._clean_text(text)
            except Exception as e:
                logger.warning(f"⚠️ pdfplumber failed: {str(e)}")
            
            # Method 2: Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    logger.info("✅ Text extracted using PyPDF2")
                    return self._clean_text(text)
            except Exception as e:
                logger.warning(f"⚠️ PyPDF2 failed: {str(e)}")
            
            # If both methods fail
                raise Exception("Could not extract text from PDF using any method")
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"❌ Failed to extract text from PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_content: bytes) -> str:
        """
        Extract text from DOCX content
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Extracted text as string
        """
        try:
            # Create a BytesIO object from the content
            docx_io = io.BytesIO(docx_content)
            
            # Load the document
            doc = Document(docx_io)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            extracted_text = '\n'.join(text_parts)
            
            if extracted_text.strip():
                logger.info("✅ Text extracted using python-docx")
                return self._clean_text(extracted_text)
            else:
                logger.warning("⚠️ No text found in DOCX file")
                return ""
                
        except Exception as e:
            logger.error(f"❌ Failed to extract text from DOCX: {str(e)}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        import re
        
        # Remove extra whitespace and normalize line breaks
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n', text)
        
        # Remove special characters that might interfere
        text = re.sub(r'[^\w\s\n.,;:()\-@]', '', text)
        
        return text.strip()
    
    def parse_resume_content(self, text: str) -> Dict[str, Any]:
        """
        Parse resume text to extract structured information using OpenAI
        
        Args:
            text: Raw resume text
            
        Returns:
            Structured resume data dictionary
        """
        try:
            # Create a prompt for OpenAI to extract structured data
            prompt = f"""
            Extract the following information from this resume text and return it in JSON format:
            
            {{
                "name": "Full name of the candidate",
                "email": "Email address",
                "phone": "Phone number",
                "skills": ["List of technical and professional skills"],
                "experience_years": "Number of years of experience (estimate if not explicit)",
                "job_titles": ["List of job titles/positions held"],
                "companies": ["List of companies worked at"],
                "education": ["Educational qualifications"],
                "certifications": ["Professional certifications"],
                "summary": "Brief professional summary (2-3 sentences)"
            }}
            
            Resume text:
            {text[:4000]}  # Limit to avoid token limits
            
            Return only valid JSON, no additional text.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a resume parsing expert. Extract information accurately and return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=1000
            )
            
            # Parse the JSON response
            import json
            parsed_data = json.loads(response.choices[0].message.content)
            
            # Ensure all fields exist with defaults
            default_data = {
                "name": "Unknown",
                "email": "",
                "phone": "",
                "skills": [],
                "experience_years": 0,
                "job_titles": [],
                "companies": [],
                "education": [],
                "certifications": [],
                "summary": ""
            }
            
            # Merge with defaults
            for key, default_value in default_data.items():
                if key not in parsed_data or not parsed_data[key]:
                    parsed_data[key] = default_value
            
            logger.info(f"✅ Resume parsed successfully for: {parsed_data.get('name', 'Unknown')}")
            return parsed_data
            
        except Exception as e:
            logger.error(f"❌ Error parsing resume content: {str(e)}")
            # Return basic structure if parsing fails
            return {
                "name": "Unknown",
                "email": "",
                "phone": "",
                "skills": [],
                "experience_years": 0,
                "job_titles": [],
                "companies": [],
                "education": [],
                "certifications": [],
                "summary": text[:500] if text else ""
            }
    
    def generate_embedding(self, text: str) -> List[float]:
        """
        Generate OpenAI embedding for the given text
        
        Args:
            text: Text to embed
            
        Returns:
            Embedding vector as list of floats
        """
        try:
            # Truncate text if too long (OpenAI has token limits)
            max_chars = 8000  # Approximately 2000 tokens
            if len(text) > max_chars:
                text = text[:max_chars]
            
            response = self.openai_client.embeddings.create(
                model=self.embedding_model,
                input=text
            )
            
            embedding = response.data[0].embedding
            logger.info(f"✅ Generated embedding with {len(embedding)} dimensions")
            
            return embedding
            
        except Exception as e:
            logger.error(f"❌ Error generating embedding: {str(e)}")
            raise
    
    def create_resume_embedding_text(self, parsed_data: Dict[str, Any], full_text: str) -> str:
        """
        Create optimized text for embedding that combines structured data with full text
        
        Args:
            parsed_data: Structured resume data
            full_text: Full resume text
            
        Returns:
            Optimized text for embedding
        """
        # Create a structured representation for better matching
        embedding_text_parts = []
        
        # Add structured information
        if parsed_data.get('name'):
            embedding_text_parts.append(f"Name: {parsed_data['name']}")
        
        if parsed_data.get('summary'):
            embedding_text_parts.append(f"Summary: {parsed_data['summary']}")
        
        if parsed_data.get('skills'):
            skills_text = ", ".join(parsed_data['skills'])
            embedding_text_parts.append(f"Skills: {skills_text}")
        
        if parsed_data.get('job_titles'):
            titles_text = ", ".join(parsed_data['job_titles'])
            embedding_text_parts.append(f"Job Titles: {titles_text}")
        
        if parsed_data.get('companies'):
            companies_text = ", ".join(parsed_data['companies'])
            embedding_text_parts.append(f"Companies: {companies_text}")
        
        if parsed_data.get('education'):
            education_text = ", ".join(parsed_data['education'])
            embedding_text_parts.append(f"Education: {education_text}")
        
        if parsed_data.get('experience_years'):
            embedding_text_parts.append(f"Experience: {parsed_data['experience_years']} years")
        
        # Combine structured data with full text
        structured_text = "\n".join(embedding_text_parts)
        combined_text = f"{structured_text}\n\nFull Resume:\n{full_text}"
        
        return combined_text
    
    def generate_resume_id(self, file_name: str, file_content: bytes) -> str:
        """
        Generate a unique ID for the resume based on filename and content
        
        Args:
            file_name: Name of the resume file
            file_content: File content bytes
            
        Returns:
            Unique resume ID
        """
        # Create hash from filename and content
        content_hash = hashlib.md5(file_content).hexdigest()
        name_hash = hashlib.md5(file_name.encode()).hexdigest()
        
        return f"resume_{name_hash[:8]}_{content_hash[:8]}"
    
    def process_resume(self, file_name: str, file_content: bytes, drive_url: str) -> Dict[str, Any]:
        """
        Complete resume processing pipeline
        
        Args:
            file_name: Name of the resume file
            file_content: PDF or DOCX content as bytes
            drive_url: Google Drive shareable URL
            
        Returns:
            Complete resume data with embedding
        """
        try:
            logger.info(f"🔄 Processing resume: {file_name}")
            
            # Generate unique ID
            resume_id = self.generate_resume_id(file_name, file_content)
            
            # Extract text based on file type
            file_extension = file_name.lower().split('.')[-1]
            if file_extension == 'pdf':
                full_text = self.extract_text_from_pdf(file_content)
            elif file_extension in ['docx', 'doc']:
                full_text = self.extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from the file")
            
            # Parse structured data
            parsed_data = self.parse_resume_content(full_text)
            
            # Create optimized text for embedding
            embedding_text = self.create_resume_embedding_text(parsed_data, full_text)
            
            # Generate embedding
            embedding = self.generate_embedding(embedding_text)
            
            # Prepare metadata for Pinecone
            metadata = {
                'name': parsed_data['name'],
                'email': parsed_data['email'],
                'phone': parsed_data['phone'],
                'skills': parsed_data['skills'][:20],  # Limit for metadata size
                'experience_years': parsed_data['experience_years'],
                'job_titles': parsed_data['job_titles'][:10],
                'companies': parsed_data['companies'][:10],
                'education': parsed_data['education'][:5],
                'drive_url': drive_url,
                'file_name': file_name,
                'summary': parsed_data['summary'][:500]  # Truncate for metadata
            }
            
            result = {
                'id': resume_id,
                'embedding': embedding,
                'metadata': metadata,
                'full_text': full_text,
                'parsed_data': parsed_data
            }
            
            logger.info(f"✅ Resume processing completed for: {parsed_data['name']}")
            return result
            
        except Exception as e:
            logger.error(f"❌ Error processing resume {file_name}: {str(e)}")
            raise
